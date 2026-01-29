import os
import re
from typing import List, Dict, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from loguru import logger
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches

class KnowledgeBaseManager:
    def __init__(self):
        """初始化知识库管理器"""
        self.vector_db_path = os.getenv("VECTOR_DB_PATH", "./data/vector_db")
        self.documents_path = os.getenv("DOCUMENTS_PATH", "./data/documents")
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "400"))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "50"))
        
        # 初始化嵌入模型
        self.embeddings = OllamaEmbeddings(
            model=os.getenv("EMBEDDING_MODEL", "dengcao/Qwen3-Embedding-0.6B:F16"),
            base_url="http://localhost:11434"
        )
        
        # 初始化向量数据库
        self.vector_store = Chroma(
            persist_directory=self.vector_db_path,
            embedding_function=self.embeddings
        )
        
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # 生成模拟数据
        self.generate_mock_data()
    
    def load_document(self, file_path: str) -> List:
        """加载不同类型的文档"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == ".pdf":
                loader = PyPDFLoader(file_path)
            elif file_extension == ".txt":
                loader = TextLoader(file_path, encoding="utf-8")
            elif file_extension in [".docx", ".doc"]:
                loader = Docx2txtLoader(file_path)
            elif file_extension in [".xlsx", ".xls"]:
                # 处理Excel文件
                import pandas as pd
                from langchain_core.documents import Document
                
                df = pd.read_excel(file_path)
                # 转换为适合检索的文本格式
                documents = []
                for _, row in df.iterrows():
                    if '分类' in df.columns and '问题' in df.columns and '答案' in df.columns:
                        content = f"分类: {row['分类']}\n问题: {row['问题']}\n答案: {row['答案']}"
                        doc = Document(page_content=content, metadata={"source": file_path})
                        documents.append(doc)
                
                logger.info(f"成功加载Excel文档: {file_path}, 行数: {len(documents)}")
                return documents
            else:
                raise ValueError(f"不支持的文件类型: {file_extension}")
            
            documents = loader.load()
            logger.info(f"成功加载文档: {file_path}, 页数: {len(documents)}")
            return documents
        except Exception as e:
            logger.error(f"加载文档失败: {str(e)}")
            raise
    
    def split_document(self, documents: List) -> List:
        """分割文档为语义块"""
        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"文档分割完成，生成 {len(chunks)} 个块")
            return chunks
        except Exception as e:
            logger.error(f"分割文档失败: {str(e)}")
            raise
    
    def add_document(self, file_path: str, document_type: str) -> bool:
        """添加文档到知识库"""
        try:
            # 加载文档
            documents = self.load_document(file_path)
            
            # 分割文档
            chunks = self.split_document(documents)
            
            # 添加元数据
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "document_type": document_type,
                    "source": os.path.basename(file_path),
                    "chunk_index": i
                })
            
            # 向量化并存储
            self.vector_store.add_documents(chunks)
            self.vector_store.persist()
            
            logger.info(f"文档成功添加到知识库: {file_path}")
            return True
        except Exception as e:
            logger.error(f"添加文档失败: {str(e)}")
            return False
    
    def generate_mock_data(self):
        """生成模拟客服数据"""
        try:
            # 创建模拟FAQ文档
            faq_content = """# 常见问题FAQ

## 退款流程
Q: 如何申请退款？
A: 您可以通过以下步骤申请退款：
1. 登录账户
2. 进入订单管理
3. 选择需要退款的订单
4. 点击"申请退款"按钮
5. 填写退款原因并提交

Q: 退款需要多长时间？
A: 退款处理时间通常为1-3个工作日，具体到账时间取决于您的支付方式。

## 价格政策
Q: 商品价格是否包含税费？
A: 是的，我们的商品价格已包含所有税费。

Q: 是否有优惠活动？
A: 我们定期举办优惠活动，请关注我们的官方网站或APP获取最新信息。

## 配送信息
Q: 配送范围包括哪些地区？
A: 我们的配送范围覆盖全国大部分地区，具体以系统显示为准。

Q: 配送时间需要多久？
A: 一般情况下，同城配送1-2天，异地配送3-5天。
"""
            
            # 创建公司介绍文档
            company_content = """# 公司介绍

## 公司简介
我们是一家专注于提供优质产品和服务的企业，成立于2010年，总部位于北京。

## 服务理念
以客户为中心，提供高效、专业的服务，不断提升客户满意度。

## 联系方式
客服热线：400-123-4567
邮箱：service@example.com
工作时间：周一至周日 9:00-21:00
"""
            
            # 创建服务条款文档
            terms_content = """# 服务条款

## 总则
本服务条款旨在规范用户与我司之间的权利义务关系。

## 用户权利
1. 有权获取我司提供的产品和服务
2. 有权对服务质量提出建议和投诉
3. 有权要求我司保护其个人信息

## 用户义务
1. 遵守国家法律法规
2. 遵守我司的各项规章制度
3. 不得利用我司服务从事违法活动

## 服务变更
我司有权根据业务需要调整服务内容，调整前会提前通知用户。
"""
            
            # 生成Excel格式的FAQ文档
            faq_file_path = os.path.join(self.documents_path, "faq.xlsx")
            wb = Workbook()
            ws = wb.active
            ws.title = "常见问题"
            
            # 添加表头
            ws.append(["分类", "问题", "答案"])
            
            # 解析FAQ内容并添加到Excel
            faq_sections = faq_content.split("\n\n")
            current_category = ""
            
            for section in faq_sections:
                if section.startswith("## "):
                    current_category = section[3:].strip()
                elif section.startswith("Q: "):
                    q_a = section.split("\nA: ")
                    if len(q_a) == 2:
                        question = q_a[0][3:].strip()
                        answer = q_a[1].strip()
                        ws.append([current_category, question, answer])
            
            wb.save(faq_file_path)
            
            # 生成Word格式的公司介绍文档
            company_file_path = os.path.join(self.documents_path, "company_intro.docx")
            doc = Document()
            doc.add_heading('公司介绍', 0)
            
            # 解析公司介绍内容并添加到Word
            company_sections = company_content.split("\n\n")
            for section in company_sections:
                if section.startswith("## "):
                    doc.add_heading(section[3:].strip(), level=1)
                else:
                    doc.add_paragraph(section.strip())
            
            doc.save(company_file_path)
            
            # 生成Word格式的服务条款文档
            terms_file_path = os.path.join(self.documents_path, "terms.docx")
            doc = Document()
            doc.add_heading('服务条款', 0)
            
            # 解析服务条款内容并添加到Word
            terms_sections = terms_content.split("\n\n")
            for section in terms_sections:
                if section.startswith("## "):
                    doc.add_heading(section[3:].strip(), level=1)
                else:
                    doc.add_paragraph(section.strip())
            
            doc.save(terms_file_path)
            
            # 添加到知识库
            mock_docs = [
                (faq_file_path, "faq"),
                (company_file_path, "company"),
                (terms_file_path, "terms")
            ]
            
            for file_path, doc_type in mock_docs:
                # 添加到知识库
                self.add_document(file_path, doc_type)
            
            logger.info("模拟数据生成完成")
        except Exception as e:
            logger.error(f"生成模拟数据失败: {str(e)}")
    
    def get_vector_store(self):
        """获取向量数据库实例"""
        return self.vector_store
    
    def clear_knowledge_base(self):
        """清空知识库"""
        try:
            # 清除向量数据库
            self.vector_store.delete_collection()
            self.vector_store = Chroma(
                persist_directory=self.vector_db_path,
                embedding_function=self.embeddings
            )
            logger.info("知识库已清空")
        except Exception as e:
            logger.error(f"清空知识库失败: {str(e)}")
            raise
